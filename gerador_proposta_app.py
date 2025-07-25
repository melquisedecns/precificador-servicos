from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import datetime

def gerar_proposta(dados):
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    def titulo(paragraph, texto):
        run = paragraph.add_run(texto)
        run.font.name = 'Calibri'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        run.font.size = Pt(12)
        run.bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Capa
    doc.add_paragraph(f"{dados['local']}, {dados['data']}.")
    doc.add_paragraph(f"{dados['codigo']} – {dados['cliente']} – {dados['titulo']}")
    doc.add_paragraph(f"\nCLIENTE: {dados['cliente']}\nLOCAL: {dados['local']}\n")
    doc.add_paragraph(f"A/C:  {dados['responsavel']}\nE-Mail: {dados['email']}\nTelefone: {dados['telefone']}")

    doc.add_paragraph("\nPROPOSTA TÉCNICA", style='Heading 1')
    doc.add_paragraph(
        f"Prezado {dados['responsavel']},\n\n"
        f"{dados['introducao']}\n\n"
        f"Permanecemos à disposição para quaisquer esclarecimentos e/ou negociações referentes a esta proposta."
    )

    doc.add_paragraph("\nAtenciosamente,\n")
    doc.add_paragraph(f"{dados['representante_1']['nome']}\nE-mail: {dados['representante_1']['email']}\nCelular: {dados['representante_1']['telefone']}")
    doc.add_paragraph(f"{dados['representante_2']['nome']}\nE-mail: {dados['representante_2']['email']}\nCelular: {dados['representante_2']['telefone']}")

    doc.add_page_break()

    doc.add_paragraph("SUMÁRIO", style='Heading 1')
    doc.add_paragraph("\n".join([
        "1. Resumo", "2. Escopo dos Serviços", "3. Materiais Fornecidos", "4. Observações",
        "5. Investimento", "6. Garantia", "7. Responsabilidades", "8. Condições", "9. Aceite"
    ]))

    doc.add_page_break()

    titulo(doc.add_paragraph(), "1. RESUMO")
    doc.add_paragraph(dados['resumo'])

    titulo(doc.add_paragraph(), "2. ESCOPO DOS SERVIÇOS")
    doc.add_paragraph("\n".join([f"✔ {item}" for item in dados['escopo']]))

    titulo(doc.add_paragraph(), "3. MATERIAIS FORNECIDOS (lista técnica)")
    doc.add_paragraph("\n".join(dados['materiais']))

    titulo(doc.add_paragraph(), "4. OBSERVAÇÕES")
    doc.add_paragraph(dados['observacoes'])

    titulo(doc.add_paragraph(), "5. INVESTIMENTO")
    investimento_txt = "\n".join([f"{k}: R$ {v:,.2f}" for k, v in dados['investimento'].items()])
    doc.add_paragraph(f"{investimento_txt}\nTOTAL: R$ {sum(dados['investimento'].values()):,.2f}")

    titulo(doc.add_paragraph(), "6. GARANTIA")
    doc.add_paragraph(dados['garantia'])

    titulo(doc.add_paragraph(), "7. RESPONSABILIDADES")
    doc.add_paragraph("Do contratado:\n" + "\n".join([f"- {item}" for item in dados['resp_contratada']]) +
                      "\n\nDo contratante:\n" + "\n".join([f"- {item}" for item in dados['resp_contratante']]))

    titulo(doc.add_paragraph(), "8. CONDIÇÕES")
    doc.add_paragraph("\n".join([f"{k.upper()}: {v}" for k, v in dados['condicoes'].items()]))

    titulo(doc.add_paragraph(), "9. ACEITE")
    doc.add_paragraph("Aceito e de acordo:\n\n______________________________________\n[Responsável]\nData: ____/____/______\n\n" + dados['assinatura'])

    nome_arquivo = f"/mnt/data/Proposta_{dados['cliente'].replace(' ', '_')}_{datetime.date.today()}.docx"
    doc.save(nome_arquivo)
    return nome_arquivo
